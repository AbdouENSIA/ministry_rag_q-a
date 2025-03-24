import os
import json
from docx import Document
import re

ROOT_DIR = "Data"

def extract_index(doc):
    """Extracts index sections from the document (assumes it's under 'الفهرس')."""
    index = {}
    capture = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if "الفهرس" in text:  # Detect where the index starts
            capture = True
            continue
        if capture and re.match(r"^\d", text):  # Match numbered sections
            match = re.search(r"(.*?)(\d+)$", text)
            if match:
                section_title = match.group(1).strip()
                page_number = int(match.group(2))
                index[section_title] = page_number
        elif capture and not text:  # Stop capturing at next empty line
            break
    return index

def extract_content_by_section(doc, index):
    """Splits content into sections based on index headings."""
    content = {}
    current_section = "Introduction"  # Default section before index
    content[current_section] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if text in index:
            current_section = text
            content[current_section] = []
        else:
            content.setdefault(current_section, []).append({"type": "paragraph", "text": text})

    # Extract tables and add them to the nearest section
    tables = extract_tables(doc)
    for table in tables:
        content.setdefault(current_section, []).append({"type": "table", "data": table})

    return content

def extract_tables(doc):
    """Extracts tables from the document and converts them into structured lists."""
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        tables.append(rows)
    return tables

def process_documents():
    """Processes all .docx files and structures them with correct section splits."""
    data = {}

    for year in os.listdir(ROOT_DIR):
        year_path = os.path.join(ROOT_DIR, year)
        if os.path.isdir(year_path):
            data[year] = {}
            for file in sorted(os.listdir(year_path)):
                if file.endswith(".docx"):
                    file_path = os.path.join(year_path, file)
                    doc_id = file.split(".")[0]
                    doc = Document(file_path)

                    index = extract_index(doc)
                    structured_content = extract_content_by_section(doc, index)

                    data[year][doc_id] = structured_content

    return data

# Process and save structured data
structured_data = process_documents()
output_path = os.path.join(ROOT_DIR, "structured_data.json")

with open(output_path, "w", encoding="utf-8") as f:
    json.dump(structured_data, f, indent=4, ensure_ascii=False)

print(f"Data successfully processed and saved to {output_path}")